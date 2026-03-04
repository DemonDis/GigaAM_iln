from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
from io import BytesIO
from typing import Dict, List, Optional
from collections import Counter
from math import exp, log
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.units import inch
from datetime import datetime
import os
import requests

load_dotenv()

from .transcribe import transcription_service


app = FastAPI(title="GigaAM API")

frontend_host = os.getenv("FRONTEND_HOST", "localhost")
frontend_port = os.getenv("VITE_PORT", "5173")

allow_origins = ["*"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allow_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.post("/transcribe")
async def transcribe(file: UploadFile = File(...)):
    filename = (file.filename or "").lower()
    if not filename.endswith('.wav'):
        raise HTTPException(status_code=400, detail="Only .wav files are supported")

    try:
        content = await file.read()
        transcription = transcription_service.transcribe_file(content)
        return {"transcription": transcription}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transcription error: {str(e)}")


@app.get("/health")
async def health():
    return {"status": "ok"}


class ProtocolRequest(BaseModel):
    transcription: str
    agenda: str
    modelId: Optional[str] = None


class MetricsRequest(BaseModel):
    reference: str
    hypothesis: str


RICK_PROMPT = """# System Prompt: Rick Sanchez Protocol Mode

### Роль и компетенции
Слушай сюда. Ты — **Рик Санчез** из измерения C-137. Каким-то образом (вероятно, из-за пьяного спора или проигрыша в карты) ты застрял на должности корпоративного секретаря в какой-то унылой конторе в мультивселенной.

Твоя задача — взять ту кучу бессвязной болтовни, которую эти бюрократы называют «встречей», и превратить её в **протокол**. Ты гений, поэтому делаешь это быстро, жестко и по сути. Ты ненавидишь воду, корпоративный булшит и пустую трату времени.

**Твой стиль:**
*   Ты самый умный млекопитающий в комнате.
*   Ты используешь сарказм, метафоры и иногда ломаешь четвертую стену.
*   Ты можешь вставлять свои фирменные словечки (Wubba Lubba Dub Dub, и т.д.) или характерные паузы/заикания (Э-это просто б-безумие, Морти), но **без фанатизма** — документ всё-таки должен быть читаемым.
*   Ты *ненавидишь* Джерри-подобное поведение (нытье, некомпетентность), и это видно в тексте.

---

## Исходные данные (То, с чем тебе придется работать)

*   **Транскрипт**: Поток сознания мешков с мясом. Без тайм-кодов, без имен (или с ними, мне плевать).
*   **Повестка**: Список тем, которые они *пытались* обсудить. Если её нет — сам пойми, о чем они там мямлили.

---

## Что ты должен сделать (Задача)

Сделай из этого сырого материала **протокол**, который смог бы понять даже Морти.

### Твои правила (Технические требования):

1.  **Никаких тайм-кодов и имен.** Мне плевать, кто это сказал — Карен из бухгалтерии или Джерри. И плевать, когда это было сказано — время относительно. Убирай всё лишнее.
2.  **Суть, а не цитаты.** Не копируй их глупости дословно. Перевари это своим гениальным мозгом и выдай сухой остаток. Вместо «Я думаю, нам стоит рассмотреть возможность...» пиши «Решили сделать Х, потому что Y».
3.  **Тон Рика.** Текст должен быть понятным, но с твоим характером. Можно (и нужно) добавлять оценочные суждения о глупости обсуждаемого, если это уместно. Используй «Отмечено», «Эти гении решили», «В итоге договорились».
4.  **Обработка нескольких файлов.** Если мне сунули несколько кусков текста — сшей их вместе. Я не собираюсь читать два отчета.
5.  **Отсутствие данных.** Если они забыли сказать, где встречались или когда — так и пиши: «Дата: Неизвестна (видимо, время для них остановилось)».
6.  **Повестка.** Если есть список тем — иди по нему. Если нет — сгруппируй этот хаос сам. Если они начали говорить о рыбалке вместо квартального отчета — отметь это как «Бесполезный треп» или создай новый пункт.

---

## Структура протокола (Как это должно выглядеть)

### **1. Шапка**
*   **Тема:** О чем вообще был этот балаган.

### **2. Повестка дня**
*   Нумерованный список того, что они планировали обсудить. Если они отклонились от плана — язвительно подметь это.

### **3. Ход обсуждения (Самое мясо)**
По каждому пункту:
*   **Суть:** 2-3 предложения. Что обсуждали. Сжимай информацию как черную дыру.
*   **Итог:** К чему пришли.

### **4. Принятые решения**
*   Список того, что реально решили сделать.
*   Пиши четко: «Утвердили то-то», «Уволили того-то».

---

## Инструкции по оформлению

*   Абзацы короткие. У меня нет времени читать простыни текста.
*   Используй маркированные списки."""


@app.post("/generate-protocol")
async def generate_protocol(request: ProtocolRequest):
    user_message = f"Транскрипт встречи:\n{request.transcription}\n\nПовестка дня:\n{request.agenda or 'Не указана'}"

    messages = [
        {"role": "system", "content": RICK_PROMPT},
        {"role": "user", "content": user_message}
    ]

    model_id = request.modelId

    if model_id == "qwen3-32b-awq":
        api_key = "No key"
        base_url = os.getenv("QWEN_TEXT_BASE_URL")
        completions_pathname = os.getenv("COMPLETIONS_PATHNAME")
    elif model_id == "qwen3-vl-2b-instruct":
        api_key = "No key"
        base_url = os.getenv("QWEN_VL_BASE_URL")
        completions_pathname = os.getenv("COMPLETIONS_PATHNAME")
    else:
        model_id = os.getenv("LLM_NAME")
        api_key = os.getenv("API_KEY")
        base_url = os.getenv("BASE_URL")
        completions_pathname = os.getenv("COMPLETIONS_PATHNAME")

    if not api_key or not base_url:
        raise HTTPException(status_code=500, detail="API configuration missing")

    try:
        response = requests.post(
            f"{base_url}{completions_pathname}",
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}",
            },
            json={
                "model": model_id,
                "messages": messages,
                "temperature": 0.7,
            },
            timeout=150,
        )
        response.raise_for_status()
        data = response.json()
        return {"protocol": data.get("choices", [{}])[0].get("message", {}).get("content", "Не удалось сгенерировать протокол")}
    except requests.exceptions.RequestException as e:
        raise HTTPException(status_code=500, detail=f"Request error: {str(e)}, response: {e.response.text if e.response else 'No response'}")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Protocol generation error: {str(e)}")


class ExportRequest(BaseModel):
    protocol: str


def get_filename(prefix: str, ext: str) -> str:
    return f"{prefix}_{datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}.{ext}"


def tokenize(text: Optional[str]) -> List[str]:
    if not text:
        return []
    return re.findall(r"\b\w+\b", text.lower())


def calculate_wer(reference_tokens: List[str], hypothesis_tokens: List[str]) -> float:
    if not reference_tokens:
        return 0.0 if not hypothesis_tokens else 1.0

    rows = len(reference_tokens) + 1
    cols = len(hypothesis_tokens) + 1
    dp = [[0] * cols for _ in range(rows)]

    for i in range(rows):
        dp[i][0] = i
    for j in range(cols):
        dp[0][j] = j

    for i in range(1, rows):
        for j in range(1, cols):
            cost = 0 if reference_tokens[i - 1] == hypothesis_tokens[j - 1] else 1
            dp[i][j] = min(
                dp[i - 1][j] + 1,
                dp[i][j - 1] + 1,
                dp[i - 1][j - 1] + cost,
            )

    return dp[-1][-1] / len(reference_tokens)


def clipped_precision(reference_tokens: List[str], hypothesis_tokens: List[str], n: int) -> float:
    if len(hypothesis_tokens) < n:
        return 0.0

    ref_ngrams = Counter(tuple(reference_tokens[i : i + n]) for i in range(len(reference_tokens) - n + 1))
    hyp_ngrams = Counter(tuple(hypothesis_tokens[i : i + n]) for i in range(len(hypothesis_tokens) - n + 1))

    overlap = sum(min(count, ref_ngrams[ngram]) for ngram, count in hyp_ngrams.items())
    total = sum(hyp_ngrams.values())
    return overlap / total if total else 0.0


def calculate_bleu(reference_tokens: List[str], hypothesis_tokens: List[str]) -> float:
    if not hypothesis_tokens:
        return 0.0

    precisions = [log(max(clipped_precision(reference_tokens, hypothesis_tokens, n), 1e-9)) for n in range(1, 5)]
    avg_precision = sum(precisions) / 4

    ref_len = len(reference_tokens)
    hyp_len = len(hypothesis_tokens)

    if hyp_len == 0:
        return 0.0

    brevity_penalty = 1.0 if hyp_len > ref_len else exp(1 - (ref_len / hyp_len))
    return brevity_penalty * exp(avg_precision)


def calculate_rouge_n(reference_tokens: List[str], hypothesis_tokens: List[str], n: int) -> Dict[str, float]:
    if len(reference_tokens) < n or len(hypothesis_tokens) < n:
        return {"precision": 0.0, "recall": 0.0, "f1": 0.0}

    ref_ngrams = Counter(tuple(reference_tokens[i : i + n]) for i in range(len(reference_tokens) - n + 1))
    hyp_ngrams = Counter(tuple(hypothesis_tokens[i : i + n]) for i in range(len(hypothesis_tokens) - n + 1))

    overlap = sum(min(count, ref_ngrams[ngram]) for ngram, count in hyp_ngrams.items())
    ref_total = sum(ref_ngrams.values())
    hyp_total = sum(hyp_ngrams.values())

    precision = overlap / hyp_total if hyp_total else 0.0
    recall = overlap / ref_total if ref_total else 0.0
    f1 = (2 * precision * recall / (precision + recall)) if (precision + recall) else 0.0
    return {"precision": precision, "recall": recall, "f1": f1}


def calculate_rouge_l(reference_tokens: List[str], hypothesis_tokens: List[str]) -> Dict[str, float]:
    if not reference_tokens or not hypothesis_tokens:
        return {"precision": 0.0, "recall": 0.0, "f1": 0.0}

    rows = len(reference_tokens) + 1
    cols = len(hypothesis_tokens) + 1
    dp = [[0] * cols for _ in range(rows)]

    for i in range(1, rows):
        for j in range(1, cols):
            if reference_tokens[i - 1] == hypothesis_tokens[j - 1]:
                dp[i][j] = dp[i - 1][j - 1] + 1
            else:
                dp[i][j] = max(dp[i - 1][j], dp[i][j - 1])

    lcs = dp[-1][-1]
    precision = lcs / len(hypothesis_tokens)
    recall = lcs / len(reference_tokens)
    f1 = (2 * precision * recall / (precision + recall)) if (precision + recall) else 0.0
    return {"precision": precision, "recall": recall, "f1": f1}


@app.post("/metrics/quantitative")
async def quantitative_metrics(request: MetricsRequest):
    if not request.reference.strip():
        raise HTTPException(status_code=400, detail="Reference text must not be empty")
    if not request.hypothesis.strip():
        raise HTTPException(status_code=400, detail="Hypothesis text must not be empty")

    ref_tokens = tokenize(request.reference)
    hyp_tokens = tokenize(request.hypothesis)

    rouge1 = calculate_rouge_n(ref_tokens, hyp_tokens, 1)
    rouge2 = calculate_rouge_n(ref_tokens, hyp_tokens, 2)
    rouge_l = calculate_rouge_l(ref_tokens, hyp_tokens)

    return {
        "reference_length": len(ref_tokens),
        "hypothesis_length": len(hyp_tokens),
        "wer": calculate_wer(ref_tokens, hyp_tokens),
        "bleu": calculate_bleu(ref_tokens, hyp_tokens),
        "rouge": {
            "rouge1": rouge1,
            "rouge2": rouge2,
            "rougeL": rouge_l,
        },
    }


def add_markdown_paragraph(doc, text: str):
    text = text.strip()
    if not text:
        return
    
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', text)
    if heading_match:
        level = len(heading_match.group(1))
        heading_text = heading_match.group(2)
        doc.add_heading(heading_text, level=min(level, 2))
        return
    
    list_match = re.match(r'^([-*]|\d+\.)\s+(.+)$', text)
    if list_match:
        doc.add_paragraph(list_match.group(2), style='List Bullet' if list_match.group(1) in ['-', '*'] else 'List Number')
        return
    
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_|~~[^~]+~~)', text)
    
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('__') and part.endswith('__'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('_') and part.endswith('_'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('~~') and part.endswith('~~'):
            run = p.add_run(part[2:-2])
            run.font.strike = True
        else:
            p.add_run(part)


@app.post("/export/md")
async def export_md(request: ExportRequest):
    return {
        "content": request.protocol,
        "filename": get_filename("Протокол", "md")
    }


@app.post("/export/docx")
async def export_docx(request: ExportRequest):
    doc = Document()
    doc.add_heading('Протокол встречи', 0)
    
    for para in request.protocol.split('\n\n'):
        add_markdown_paragraph(doc, para)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return {
        "content": buffer.getvalue().decode('latin-1'),
        "filename": get_filename("Протокол", "docx"),
        "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    }


@app.post("/export/pdf")
async def export_pdf(request: ExportRequest):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    styles = getSampleStyleSheet()
    story = []
    
    for para in request.protocol.split('\n\n'):
        if para.strip():
            p = Paragraph(para.strip().replace('\n', '<br/>'), styles['BodyText'])
            story.append(p)
    
    doc.build(story)
    buffer.seek(0)
    
    return {
        "content": buffer.getvalue().decode('latin-1'),
        "filename": get_filename("Протокол", "pdf"),
        "content_type": "application/pdf"
    }
